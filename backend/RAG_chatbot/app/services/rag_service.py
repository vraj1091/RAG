# Enhanced RAG Service Implementation
# This file contains the complete enhanced backend solution

import os
import json
import logging
from typing import List, Dict, Any, Optional
import numpy as np
from datetime import datetime
import requests
import pandas as pd
import gdown
from urllib.parse import urlparse, parse_qs
import tempfile
import io

# Enhanced imports for new features
import openpyxl
from openpyxl import load_workbook
import trafilatura
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from webdriver_manager.chrome import ChromeDriverManager
import time
import re

# Google Drive API imports
try:
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    from google.auth.transport.requests import Request
    from google.oauth2.service_account import Credentials
    GOOGLE_DRIVE_AVAILABLE = True
except ImportError:
    GOOGLE_DRIVE_AVAILABLE = False
    logging.warning("Google Drive API not available. Install google-api-python-client")

# Web scraping imports
from bs4 import BeautifulSoup
import html2text

# RAG and LLM imports
import chromadb
from chromadb.config import Settings as ChromaSettings
from sentence_transformers import SentenceTransformer
from langchain.text_splitter import RecursiveCharacterTextSplitter

# CRITICAL FIX: Use DeepSeek instead of Gemini for privacy
from app.services.deepseek_service import deepseek_service

# Document processing imports
import PyPDF2
import docx
from PIL import Image
import pytesseract
import cv2

from app.core.config import settings
from app.models.models import Document, DocumentChunk
from sqlalchemy.orm import Session

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class GoogleDriveService:
    """Enhanced Google Drive integration service"""
    
    def __init__(self, credentials_path: str = None):
        """Initialize Google Drive service with credentials"""
        self.service = None
        if credentials_path and GOOGLE_DRIVE_AVAILABLE:
            try:
                credentials = Credentials.from_service_account_file(credentials_path)
                self.service = build('drive', 'v3', credentials=credentials)
                logger.info("✅ Google Drive API initialized successfully")
            except Exception as e:
                logger.error(f"❌ Failed to initialize Google Drive API: {e}")
        elif not GOOGLE_DRIVE_AVAILABLE:
            logger.warning("⚠️ Google Drive API dependencies not installed")
    
    def extract_file_id_from_url(self, url: str) -> Optional[str]:
        """Extract file ID from various Google Drive URL formats"""
        # Pattern 1: /file/d/{file_id}/
        match = re.search(r'/file/d/([a-zA-Z0-9-_]+)', url)
        if match:
            return match.group(1)
        
        # Pattern 2: ?id={file_id}
        match = re.search(r'[?&]id=([a-zA-Z0-9-_]+)', url)
        if match:
            return match.group(1)
        
        # Pattern 3: /folders/{folder_id}
        match = re.search(r'/folders/([a-zA-Z0-9-_]+)', url)
        if match:
            return match.group(1)
        
        return None
    
    def get_file_metadata(self, file_id: str) -> Dict[str, Any]:
        """Get file metadata from Google Drive"""
        if not self.service:
            return {}
        
        try:
            file_metadata = self.service.files().get(
                fileId=file_id, 
                fields='id,name,mimeType,size,parents'
            ).execute()
            return file_metadata
        except Exception as e:
            logger.error(f"❌ Failed to get file metadata for {file_id}: {e}")
            return {}
    
    def download_file_content(self, file_id: str, mime_type: str) -> Optional[bytes]:
        """Download file content from Google Drive"""
        if not self.service:
            return None
        
        try:
            # Handle Google Docs, Sheets, etc. (export as different formats)
            if mime_type == 'application/vnd.google-apps.document':
                request = self.service.files().export_media(
                    fileId=file_id, 
                    mimeType='text/plain'
                )
            elif mime_type == 'application/vnd.google-apps.spreadsheet':
                request = self.service.files().export_media(
                    fileId=file_id, 
                    mimeType='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
                )
            else:
                request = self.service.files().get_media(fileId=file_id)
            
            file_io = io.BytesIO()
            downloader = MediaIoBaseDownload(file_io, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            return file_io.getvalue()
            
        except Exception as e:
            logger.error(f"❌ Failed to download file {file_id}: {e}")
            return None
    
    def list_folder_files(self, folder_id: str) -> List[Dict[str, Any]]:
        """List all files in a Google Drive folder"""
        if not self.service:
            return []
        
        try:
            results = self.service.files().list(
                q=f"'{folder_id}' in parents and trashed=false",
                fields='files(id,name,mimeType,size)',
                pageSize=100
            ).execute()
            
            return results.get('files', [])
            
        except Exception as e:
            logger.error(f"❌ Failed to list folder contents for {folder_id}: {e}")
            return []
class EnhancedRAGService:
    """
    Enhanced RAG service with DeepSeek R1, Google Drive, advanced web scraping,
    and improved Excel processing capabilities
    """
    
    def __init__(self, db: Session, user_id: int):
        """Initialize Enhanced RAG service"""
        self.db = db
        self.user_id = user_id
        
        # Initialize embedding model
        try:
            self.embedding_model = SentenceTransformer(settings.embedding_model)
            logger.info(f"✅ Loaded embedding model: {settings.embedding_model}")
        except Exception as e:
            logger.error(f"❌ Failed to load embedding model: {e}")
            self.embedding_model = None
        
        # Initialize ChromaDB with persistent storage
        try:
            chroma_client = chromadb.PersistentClient(
                path=settings.vector_store_path,
                settings=ChromaSettings(anonymized_telemetry=False)
            )
            self.collection = chroma_client.get_or_create_collection(
                name=f"user_{user_id}_documents"
            )
            logger.info(f"✅ ChromaDB collection initialized for user {user_id}")
        except Exception as e:
            logger.error(f"❌ Failed to initialize ChromaDB: {e}")
            self.collection = None
        
        # CRITICAL FIX: Use DeepSeek instead of Gemini
        self.gemini_model = deepseek_service
        logger.info("✅ DeepSeek R1 initialized (local & private)")
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
        )
        
        # Initialize Google Drive service if enabled
        self.google_drive_service = None
        if settings.enable_google_drive and settings.google_drive_credentials_path:
            self.google_drive_service = GoogleDriveService(
                credentials_path=settings.google_drive_credentials_path
            )
        
        # Initialize HTML to text converter
        self.html_converter = html2text.HTML2Text()
        self.html_converter.ignore_links = False
        self.html_converter.ignore_images = True
    
    def process_document_sources(
        self, 
        sources: List[str], 
        source_type: str = "file"
    ) -> Dict[str, Any]:
        """
        Enhanced method to process multiple document sources
        Supports: local files, URLs, Google Drive files/folders
        """
        results = {
            "processed": [],
            "failed": [],
            "total_chunks": 0,
            "errors": []
        }
        
        for source in sources:
            try:
                if source_type == "file":
                    # Local file processing
                    result = self._process_local_file(source)
                elif source_type == "url":
                    # Enhanced web scraping
                    result = self._process_url_enhanced(source)
                elif source_type == "google_drive":
                    # Google Drive file/folder
                    result = self._process_google_drive_source(source)
                else:
                    result = {"error": f"Unknown source type: {source_type}"}
                
                if "error" in result:
                    results["failed"].append({
                        "source": source,
                        "error": result["error"]
                    })
                    results["errors"].append(result["error"])
                else:
                    results["processed"].append({
                        "source": source,
                        "chunks": result.get("chunks", 0),
                        "doc_id": result.get("doc_id")
                    })
                    results["total_chunks"] += result.get("chunks", 0)
                    
            except Exception as e:
                logger.error(f"❌ Error processing source {source}: {e}")
                results["failed"].append({
                    "source": source,
                    "error": str(e)
                })
                results["errors"].append(str(e))
        
        return results
    
    def _process_local_file(self, file_path: str) -> Dict[str, Any]:
        """Process a local file"""
        try:
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Extract text based on file type
            if file_ext == '.pdf':
                text = self._extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                text = self._extract_text_from_docx(file_path)
            elif file_ext in ['.xlsx', '.xls', '.csv']:
                text = self._extract_text_from_excel_enhanced(file_path)
            elif file_ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_ext in ['.png', '.jpg', '.jpeg']:
                text = self._extract_text_from_image_ocr(file_path)
            else:
                return {"error": f"Unsupported file type: {file_ext}"}
            
            if not text or len(text.strip()) < 50:
                return {"error": "Extracted text too short or empty"}
            
            # Store document and create chunks
            doc_id = self._store_document(
                filename=os.path.basename(file_path),
                content=text,
                file_type=file_ext,
                source_type="file"
            )
            
            chunks = self._create_and_store_chunks(text, doc_id)
            
            return {
                "doc_id": doc_id,
                "chunks": len(chunks),
                "text_length": len(text)
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing local file {file_path}: {e}")
            return {"error": str(e)}
    
    def _process_google_drive_source(self, drive_url: str) -> Dict[str, Any]:
        """Process Google Drive file or folder"""
        if not self.google_drive_service or not self.google_drive_service.service:
            return {"error": "Google Drive service not initialized"}
        
        try:
            file_id = self.google_drive_service.extract_file_id_from_url(drive_url)
            if not file_id:
                return {"error": "Invalid Google Drive URL"}
            
            # Get file metadata
            metadata = self.google_drive_service.get_file_metadata(file_id)
            if not metadata:
                return {"error": "Failed to get file metadata"}
            
            mime_type = metadata.get('mimeType', '')
            
            # Check if it's a folder
            if mime_type == 'application/vnd.google-apps.folder':
                return self._process_google_drive_folder(file_id, metadata['name'])
            else:
                return self._process_google_drive_file(file_id, metadata)
                
        except Exception as e:
            logger.error(f"❌ Error processing Google Drive source: {e}")
            return {"error": str(e)}
    
    def _process_google_drive_file(
        self, 
        file_id: str, 
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Process a single Google Drive file"""
        try:
            file_name = metadata.get('name', 'Unknown')
            mime_type = metadata.get('mimeType', '')
            
            # Download file content
            content_bytes = self.google_drive_service.download_file_content(
                file_id, 
                mime_type
            )
            
            if not content_bytes:
                return {"error": f"Failed to download file: {file_name}"}
            
            # Extract text based on mime type
            if 'text/plain' in mime_type or 'vnd.google-apps.document' in mime_type:
                text = content_bytes.decode('utf-8')
            elif 'application/pdf' in mime_type:
                # Save temporarily and process
                with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                    tmp.write(content_bytes)
                    tmp_path = tmp.name
                text = self._extract_text_from_pdf(tmp_path)
                os.unlink(tmp_path)
            elif 'spreadsheet' in mime_type:
                # Save temporarily and process as Excel
                with tempfile.NamedTemporaryFile(suffix='.xlsx', delete=False) as tmp:
                    tmp.write(content_bytes)
                    tmp_path = tmp.name
                text = self._extract_text_from_excel_enhanced(tmp_path)
                os.unlink(tmp_path)
            else:
                return {"error": f"Unsupported mime type: {mime_type}"}
            
            if not text or len(text.strip()) < 50:
                return {"error": "Extracted text too short"}
            
            # Store document
            doc_id = self._store_document(
                filename=file_name,
                content=text,
                file_type="google_drive",
                source_type="google_drive",
                source_url=f"https://drive.google.com/file/d/{file_id}"
            )
            
            chunks = self._create_and_store_chunks(text, doc_id)
            
            return {
                "doc_id": doc_id,
                "chunks": len(chunks),
                "text_length": len(text)
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing Google Drive file: {e}")
            return {"error": str(e)}
    
    def _process_google_drive_folder(
        self, 
        folder_id: str, 
        folder_name: str
    ) -> Dict[str, Any]:
        """Process all files in a Google Drive folder"""
        try:
            files = self.google_drive_service.list_folder_files(folder_id)
            
            if not files:
                return {"error": "No files found in folder"}
            
            total_chunks = 0
            processed_files = []
            failed_files = []
            
            # Limit number of files to process
            max_files = settings.max_google_drive_folder_files
            files = files[:max_files]
            
            for file_metadata in files:
                result = self._process_google_drive_file(
                    file_metadata['id'], 
                    file_metadata
                )
                
                if "error" in result:
                    failed_files.append({
                        "name": file_metadata.get('name'),
                        "error": result["error"]
                    })
                else:
                    processed_files.append({
                        "name": file_metadata.get('name'),
                        "chunks": result.get("chunks", 0)
                    })
                    total_chunks += result.get("chunks", 0)
            
            return {
                "folder_name": folder_name,
                "total_files": len(files),
                "processed": len(processed_files),
                "failed": len(failed_files),
                "total_chunks": total_chunks,
                "processed_files": processed_files,
                "failed_files": failed_files
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing Google Drive folder: {e}")
            return {"error": str(e)}
    def _process_url_enhanced(self, url: str) -> Dict[str, Any]:
        """Enhanced URL processing with Selenium fallback"""
        try:
            # First try with trafilatura (fast)
            if settings.enable_enhanced_web_scraping:
                text = self._scrape_with_trafilatura(url)
                
                # If trafilatura fails or returns short content, try Selenium
                if not text or len(text.strip()) < 100:
                    if settings.enable_selenium:
                        text = self._scrape_with_selenium(url)
                    else:
                        text = self._scrape_with_requests(url)
            else:
                text = self._scrape_with_requests(url)
            
            if not text or len(text.strip()) < 50:
                return {"error": "Failed to extract meaningful content from URL"}
            
            # Store document
            doc_id = self._store_document(
                filename=url,
                content=text,
                file_type="url",
                source_type="url",
                source_url=url
            )
            
            chunks = self._create_and_store_chunks(text, doc_id)
            
            return {
                "doc_id": doc_id,
                "chunks": len(chunks),
                "text_length": len(text)
            }
            
        except Exception as e:
            logger.error(f"❌ Error processing URL {url}: {e}")
            return {"error": str(e)}
    
    def _scrape_with_trafilatura(self, url: str) -> Optional[str]:
        """Scrape URL using trafilatura (fast and clean)"""
        try:
            downloaded = trafilatura.fetch_url(url)
            if downloaded:
                text = trafilatura.extract(downloaded)
                return text
        except Exception as e:
            logger.warning(f"⚠️ Trafilatura scraping failed for {url}: {e}")
        return None
    
    def _scrape_with_selenium(self, url: str) -> Optional[str]:
        """Scrape URL using Selenium (for dynamic content)"""
        try:
            options = Options()
            if settings.selenium_headless:
                options.add_argument('--headless')
            options.add_argument('--no-sandbox')
            options.add_argument('--disable-dev-shm-usage')
            options.add_argument(f'user-agent={settings.user_agent}')
            
            driver = webdriver.Chrome(
                service=webdriver.chrome.service.Service(ChromeDriverManager().install()),
                options=options
            )
            
            driver.get(url)
            
            # Wait for page to load
            WebDriverWait(driver, settings.selenium_wait_time).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )
            
            # Get page source
            page_source = driver.page_source
            driver.quit()
            
            # Parse with BeautifulSoup
            soup = BeautifulSoup(page_source, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            # Get text
            text = soup.get_text(separator='\n', strip=True)
            
            return text
            
        except Exception as e:
            logger.error(f"❌ Selenium scraping failed for {url}: {e}")
            return None
    
    def _scrape_with_requests(self, url: str) -> Optional[str]:
        """Basic scraping using requests and BeautifulSoup"""
        try:
            headers = {
                'User-Agent': settings.user_agent
            }
            response = requests.get(
                url, 
                headers=headers, 
                timeout=settings.web_scraping_timeout
            )
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove unwanted elements
            for script in soup(["script", "style", "nav", "footer", "header"]):
                script.decompose()
            
            text = soup.get_text(separator='\n', strip=True)
            
            return text
            
        except Exception as e:
            logger.error(f"❌ Requests scraping failed for {url}: {e}")
            return None
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            # If text extraction failed, try OCR
            if len(text.strip()) < 100:
                logger.info("PDF text extraction yielded little content, trying OCR...")
                text = self._extract_text_from_pdf_with_ocr(pdf_path)
            
            return text.strip()
        except Exception as e:
            logger.error(f"❌ Error extracting text from PDF: {e}")
            return ""
    
    def _extract_text_from_pdf_with_ocr(self, pdf_path: str) -> str:
        """Extract text from PDF using OCR"""
        try:
            from pdf2image import convert_from_path
            
            images = convert_from_path(
                pdf_path,
                poppler_path=settings.poppler_path if os.name == 'nt' else None
            )
            
            text = ""
            for i, image in enumerate(images):
                # Convert PIL image to numpy array for OpenCV
                img_array = np.array(image)
                
                # Preprocess image
                processed_img = self._preprocess_image_for_ocr(img_array)
                
                # Extract text
                page_text = pytesseract.image_to_string(processed_img)
                text += f"\n--- Page {i+1} ---\n{page_text}\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"❌ Error in PDF OCR: {e}")
            return ""
    
    def _extract_text_from_docx(self, docx_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_path)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += "\n" + row_text
            
            return text.strip()
        except Exception as e:
            logger.error(f"❌ Error extracting text from DOCX: {e}")
            return ""
    
    def _extract_text_from_excel_enhanced(self, excel_path: str) -> str:
        """Enhanced Excel text extraction with better formatting"""
        try:
            # Try reading with pandas first
            try:
                df = pd.read_excel(excel_path, sheet_name=None)  # Read all sheets
                
                text = ""
                for sheet_name, sheet_df in df.items():
                    text += f"\n\n=== Sheet: {sheet_name} ===\n\n"
                    
                    # Convert DataFrame to formatted text
                    text += sheet_df.to_string(index=False)
                    text += "\n"
                
                return text.strip()
            except Exception as e:
                logger.warning(f"⚠️ Pandas Excel read failed, trying openpyxl: {e}")
                
                # Fallback to openpyxl for more detailed extraction
                workbook = load_workbook(excel_path, data_only=True)
                text = ""
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"\n\n=== Sheet: {sheet_name} ===\n\n"
                    
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip():
                            text += row_text + "\n"
                
                return text.strip()
                
        except Exception as e:
            logger.error(f"❌ Error extracting text from Excel: {e}")
            return ""
    
    def _extract_text_from_image_ocr(self, image_path: str) -> str:
        """Extract text from image using OCR"""
        try:
            # Load image
            image = cv2.imread(image_path)
            
            if image is None:
                logger.error(f"❌ Failed to load image: {image_path}")
                return ""
            
            # Preprocess image
            processed_image = self._preprocess_image_for_ocr(image)
            
            # Perform OCR
            text = pytesseract.image_to_string(processed_image)
            
            return text.strip()
        except Exception as e:
            logger.error(f"❌ Error in image OCR: {e}")
            return ""
    
    def _preprocess_image_for_ocr(self, image: np.ndarray) -> np.ndarray:
        """Preprocess image to improve OCR accuracy"""
        try:
            # Convert to grayscale
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # Apply denoising
            denoised = cv2.fastNlMeansDenoising(gray)
            
            # Apply adaptive thresholding
            threshold = cv2.adaptiveThreshold(
                denoised,
                255,
                cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
                cv2.THRESH_BINARY,
                11,
                2
            )
            
            return threshold
        except Exception as e:
            logger.warning(f"⚠️ Image preprocessing failed, using original: {e}")
            return image
    def _store_document(
        self,
        filename: str,
        content: str,
        file_type: str,
        source_type: str = "file",
        source_url: Optional[str] = None
    ) -> int:
        """Store document in database"""
        try:
            document = Document(
                user_id=self.user_id,
                filename=filename,
                original_filename=filename,
                file_type=file_type,
                file_path=source_url or filename,
                content=content,
                created_at=datetime.utcnow()
            )
            
            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)
            
            logger.info(f"✅ Stored document: {filename} (ID: {document.id})")
            return document.id
            
        except Exception as e:
            logger.error(f"❌ Error storing document: {e}")
            self.db.rollback()
            raise
    
    def _create_and_store_chunks(self, text: str, document_id: int) -> List[str]:
        """Create and store text chunks in vector database"""
        try:
            # Split text into chunks
            chunks = self.text_splitter.split_text(text)
            
            # Limit number of chunks
            if len(chunks) > settings.max_chunks_per_document:
                logger.warning(
                    f"⚠️ Document has {len(chunks)} chunks, limiting to {settings.max_chunks_per_document}"
                )
                chunks = chunks[:settings.max_chunks_per_document]
            
            if not self.embedding_model or not self.collection:
                logger.error("❌ Embedding model or collection not initialized")
                return []
            
            # Generate embeddings
            embeddings = self.embedding_model.encode(chunks).tolist()
            
            # Store in ChromaDB
            chunk_ids = [f"doc_{document_id}_chunk_{i}" for i in range(len(chunks))]
            
            self.collection.add(
                ids=chunk_ids,
                embeddings=embeddings,
                documents=chunks,
                metadatas=[
                    {
                        "document_id": document_id,
                        "chunk_index": i,
                        "user_id": self.user_id
                    }
                    for i in range(len(chunks))
                ]
            )
            
            # Store chunks in database
            for i, chunk_text in enumerate(chunks):
                chunk = DocumentChunk(
                    document_id=document_id,
                    chunk_index=i,
                    content=chunk_text,
                    embedding=json.dumps(embeddings[i])
                )
                self.db.add(chunk)
            
            self.db.commit()
            
            logger.info(f"✅ Created and stored {len(chunks)} chunks for document {document_id}")
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Error creating chunks: {e}")
            self.db.rollback()
            return []
    
    def query_documents(self, query: str, top_k: int = 5) -> Dict[str, Any]:
        """
        Query documents using RAG with DeepSeek R1
        """
        try:
            if not self.embedding_model or not self.collection:
                return {
                    "answer": "❌ RAG system not properly initialized",
                    "sources": []
                }
            
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0].tolist()
            
            # Search in ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                where={"user_id": self.user_id}
            )
            
            if not results['documents'] or len(results['documents'][0]) == 0:
                return {
                    "answer": "I don't have any relevant information in my knowledge base to answer this question.",
                    "sources": []
                }
            
            # Combine retrieved chunks as context
            context_chunks = results['documents'][0]
            context = "\n\n".join(context_chunks)
            
            # Generate response with DeepSeek R1 (using local model)
            if self.gemini_model:
                prompt = f"""Based on the following context, please answer the question.

Context:
{context}

Question: {query}

Answer:"""
                
                try:
                    response = self.gemini_model.generate_content(prompt)
                    answer = response.text
                except Exception as e:
                    logger.error(f"❌ DeepSeek generation error: {e}")
                    answer = f"Error generating response: {str(e)}"
            else:
                answer = "AI model not available"
            
            # Get source documents
            metadatas = results['metadatas'][0] if results['metadatas'] else []
            source_doc_ids = list(set([m.get('document_id') for m in metadatas if 'document_id' in m]))
            
            sources = []
            for doc_id in source_doc_ids:
                doc = self.db.query(Document).filter(
                    Document.id == doc_id,
                    Document.user_id == self.user_id
                ).first()
                
                if doc:
                    sources.append({
                        "id": doc.id,
                        "filename": doc.filename,
                        "file_type": doc.file_type
                    })
            
            return {
                "answer": answer,
                "sources": sources,
                "context_chunks": len(context_chunks)
            }
            
        except Exception as e:
            logger.error(f"❌ Error in query_documents: {e}")
            return {
                "answer": f"Error processing query: {str(e)}",
                "sources": []
            }
    
    def get_all_documents(self) -> List[Dict[str, Any]]:
        """Get all documents for the user"""
        try:
            documents = self.db.query(Document).filter(
                Document.user_id == self.user_id
            ).all()
            
            return [
                {
                    "id": doc.id,
                    "filename": doc.filename,
                    "file_type": doc.file_type,
                    "created_at": doc.created_at.isoformat() if doc.created_at else None,
                    "content_length": len(doc.content) if doc.content else 0
                }
                for doc in documents
            ]
        except Exception as e:
            logger.error(f"❌ Error getting documents: {e}")
            return []
    
    def delete_document(self, document_id: int) -> bool:
        """Delete a document and its chunks"""
        try:
            # Get document
            document = self.db.query(Document).filter(
                Document.id == document_id,
                Document.user_id == self.user_id
            ).first()
            
            if not document:
                logger.warning(f"⚠️ Document {document_id} not found")
                return False
            
            # Delete chunks from ChromaDB
            if self.collection:
                try:
                    # Get all chunk IDs for this document
                    chunk_ids = [
                        f"doc_{document_id}_chunk_{i}"
                        for i in range(settings.max_chunks_per_document)
                    ]
                    
                    # Delete from ChromaDB (it will ignore non-existent IDs)
                    self.collection.delete(ids=chunk_ids)
                except Exception as e:
                    logger.warning(f"⚠️ Error deleting chunks from ChromaDB: {e}")
            
            # Delete chunks from database
            self.db.query(DocumentChunk).filter(
                DocumentChunk.document_id == document_id
            ).delete()
            
            # Delete document
            self.db.delete(document)
            self.db.commit()
            
            logger.info(f"✅ Deleted document {document_id}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error deleting document: {e}")
            self.db.rollback()
            return False
    
    def delete_all_documents(self) -> Dict[str, Any]:
        """Delete all documents for the user"""
        try:
            # Get all documents
            documents = self.db.query(Document).filter(
                Document.user_id == self.user_id
            ).all()
            
            deleted_count = 0
            for document in documents:
                if self.delete_document(document.id):
                    deleted_count += 1
            
            # Also clear the ChromaDB collection
            if self.collection:
                try:
                    # Delete all entries for this user
                    self.collection.delete(
                        where={"user_id": self.user_id}
                    )
                except Exception as e:
                    logger.warning(f"⚠️ Error clearing ChromaDB collection: {e}")
            
            return {
                "deleted": deleted_count,
                "total": len(documents)
            }
            
        except Exception as e:
            logger.error(f"❌ Error deleting all documents: {e}")
            return {
                "deleted": 0,
                "total": 0,
                "error": str(e)
            }
    
    def get_document_stats(self) -> Dict[str, Any]:
        """Get statistics about user's documents"""
        try:
            documents = self.db.query(Document).filter(
                Document.user_id == self.user_id
            ).all()
            
            total_documents = len(documents)
            total_chunks = self.db.query(DocumentChunk).join(
                Document
            ).filter(
                Document.user_id == self.user_id
            ).count()
            
            file_types = {}
            for doc in documents:
                file_type = doc.file_type
                file_types[file_type] = file_types.get(file_type, 0) + 1
            
            return {
                "total_documents": total_documents,
                "total_chunks": total_chunks,
                "file_types": file_types,
                "avg_chunks_per_doc": round(total_chunks / total_documents, 2) if total_documents > 0 else 0
            }
            
        except Exception as e:
            logger.error(f"❌ Error getting document stats: {e}")
            return {
                "total_documents": 0,
                "total_chunks": 0,
                "file_types": {},
                "avg_chunks_per_doc": 0
            }
    async def generate_rag_response(self, query: str, user_id: int) -> dict:
        """
        Generate RAG response - compatibility method for chat.py
        Maps to query_documents() method
        """
        try:
            # Use the existing query_documents method
            result = self.query_documents(query, top_k=5)
            
            # Format response to match expected structure
            return {
                "response": result.get("answer", ""),
                "sources": result.get("sources", []),
                "context_used": len(result.get("sources", [])) > 0,
                "chunks_retrieved": result.get("context_chunks", 0)
            }
        except Exception as e:
            logger.error(f"Error in generate_rag_response: {e}")
            return {
                "response": f"Error: {str(e)}",
                "sources": [],
                "context_used": False,
                "chunks_retrieved": 0
            }

    def search_documents(self, search_query: str) -> List[Dict[str, Any]]:
        """Search documents by filename or content"""
        try:
            documents = self.db.query(Document).filter(
                Document.user_id == self.user_id,
                (
                    Document.filename.contains(search_query) |
                    Document.content.contains(search_query)
                )
            ).all()
            
            return [
                {
                    "id": doc.id,
                    "filename": doc.filename,
                    "file_type": doc.file_type,
                    "created_at": doc.created_at.isoformat() if doc.created_at else None,
                    "content_preview": doc.content[:200] + "..." if doc.content and len(doc.content) > 200 else doc.content
                }
                for doc in documents
            ]
        except Exception as e:
            logger.error(f"❌ Error searching documents: {e}")
            return []
    
    def get_document_content(self, document_id: int) -> Optional[Dict[str, Any]]:
        """Get full document content"""
        try:
            document = self.db.query(Document).filter(
                Document.id == document_id,
                Document.user_id == self.user_id
            ).first()
            
            if not document:
                return None
            
            return {
                "id": document.id,
                "filename": document.filename,
                "file_type": document.file_type,
                "content": document.content,
                "created_at": document.created_at.isoformat() if document.created_at else None
            }
        except Exception as e:
            logger.error(f"❌ Error getting document content: {e}")
            return None
    
    def export_knowledge_base(self, export_format: str = "json") -> Optional[str]:
        """Export entire knowledge base"""
        try:
            documents = self.db.query(Document).filter(
                Document.user_id == self.user_id
            ).all()
            
            if export_format == "json":
                export_data = {
                    "user_id": self.user_id,
                    "export_date": datetime.utcnow().isoformat(),
                    "total_documents": len(documents),
                    "documents": [
                        {
                            "id": doc.id,
                            "filename": doc.filename,
                            "file_type": doc.file_type,
                            "content": doc.content,
                            "created_at": doc.created_at.isoformat() if doc.created_at else None
                        }
                        for doc in documents
                    ]
                }
                
                return json.dumps(export_data, indent=2)
            
            elif export_format == "txt":
                text_export = f"Knowledge Base Export\n"
                text_export += f"User ID: {self.user_id}\n"
                text_export += f"Export Date: {datetime.utcnow().isoformat()}\n"
                text_export += f"Total Documents: {len(documents)}\n"
                text_export += "=" * 80 + "\n\n"
                
                for doc in documents:
                    text_export += f"\n{'=' * 80}\n"
                    text_export += f"Document: {doc.filename}\n"
                    text_export += f"Type: {doc.file_type}\n"
                    text_export += f"Created: {doc.created_at.isoformat() if doc.created_at else 'Unknown'}\n"
                    text_export += f"{'=' * 80}\n\n"
                    text_export += doc.content + "\n"
                
                return text_export
            
            else:
                logger.error(f"❌ Unsupported export format: {export_format}")
                return None
                
        except Exception as e:
            logger.error(f"❌ Error exporting knowledge base: {e}")
            return None
    
    def get_similar_chunks(
        self, 
        query: str, 
        top_k: int = 10
    ) -> List[Dict[str, Any]]:
        """Get similar chunks without generating an answer"""
        try:
            if not self.embedding_model or not self.collection:
                return []
            
            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])[0].tolist()
            
            # Search in ChromaDB
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                where={"user_id": self.user_id}
            )
            
            if not results['documents'] or len(results['documents'][0]) == 0:
                return []
            
            chunks = []
            for i, chunk_text in enumerate(results['documents'][0]):
                metadata = results['metadatas'][0][i] if results['metadatas'] else {}
                distance = results['distances'][0][i] if results.get('distances') else None
                
                chunks.append({
                    "content": chunk_text,
                    "document_id": metadata.get('document_id'),
                    "chunk_index": metadata.get('chunk_index'),
                    "distance": distance
                })
            
            return chunks
            
        except Exception as e:
            logger.error(f"❌ Error getting similar chunks: {e}")
            return []
    
    def batch_process_urls(self, urls: List[str]) -> Dict[str, Any]:
        """Batch process multiple URLs"""
        return self.process_document_sources(urls, source_type="url")
    
    def batch_process_files(self, file_paths: List[str]) -> Dict[str, Any]:
        """Batch process multiple files"""
        return self.process_document_sources(file_paths, source_type="file")
    
    def update_document_metadata(
        self, 
        document_id: int, 
        metadata: Dict[str, Any]
    ) -> bool:
        """Update document metadata"""
        try:
            document = self.db.query(Document).filter(
                Document.id == document_id,
                Document.user_id == self.user_id
            ).first()
            
            if not document:
                return False
            
            # Update allowed fields
            if 'filename' in metadata:
                document.filename = metadata['filename']
            if 'original_filename' in metadata:
                document.original_filename = metadata['original_filename']
            
            self.db.commit()
            return True
            
        except Exception as e:
            logger.error(f"❌ Error updating document metadata: {e}")
            self.db.rollback()
            return False
    
    def get_system_info(self) -> Dict[str, Any]:
        """Get system information and capabilities"""
        return {
            "embedding_model": settings.embedding_model,
            "ai_model": "DeepSeek R1 8B (Local & Private)",
            "vector_db": settings.vector_db_type,
            "chunk_size": settings.chunk_size,
            "chunk_overlap": settings.chunk_overlap,
            "google_drive_enabled": self.google_drive_service is not None,
            "selenium_enabled": settings.enable_selenium,
            "enhanced_web_scraping": settings.enable_enhanced_web_scraping,
            "supported_file_types": settings.allowed_file_types_list,
            "max_file_size_mb": settings.max_file_size_mb,
            "ollama_base_url": settings.ollama_base_url,
            "ollama_model": settings.ollama_model_name,
            "privacy_mode": "FULL - All data stays local"
        }
    
    def health_check(self) -> Dict[str, Any]:
        """Check health of RAG system components"""
        health = {
            "status": "healthy",
            "components": {}
        }
        
        # Check embedding model
        health["components"]["embedding_model"] = {
            "status": "ok" if self.embedding_model else "error",
            "name": settings.embedding_model
        }
        
        # Check vector database
        health["components"]["vector_db"] = {
            "status": "ok" if self.collection else "error",
            "type": settings.vector_db_type
        }
        
        # Check DeepSeek model
        health["components"]["ai_model"] = {
            "status": "ok" if self.gemini_model else "error",
            "name": "DeepSeek R1 8B",
            "location": "Local (Ollama)"
        }
        
        # Check Google Drive
        health["components"]["google_drive"] = {
            "status": "ok" if self.google_drive_service and self.google_drive_service.service else "disabled",
            "enabled": settings.enable_google_drive
        }
        
        # Overall status
        if any(
            comp["status"] == "error" 
            for comp in health["components"].values()
        ):
            health["status"] = "degraded"
        
        return health
    
    def reindex_documents(self) -> Dict[str, Any]:
        """Reindex all documents (rebuild vector database)"""
        try:
            # Get all documents
            documents = self.db.query(Document).filter(
                Document.user_id == self.user_id
            ).all()
            
            if not documents:
                return {
                    "status": "no_documents",
                    "message": "No documents to reindex"
                }
            
            # Clear existing chunks
            if self.collection:
                try:
                    self.collection.delete(where={"user_id": self.user_id})
                except Exception as e:
                    logger.warning(f"⚠️ Error clearing collection: {e}")
            
            # Delete chunks from database
            for doc in documents:
                self.db.query(DocumentChunk).filter(
                    DocumentChunk.document_id == doc.id
                ).delete()
            
            self.db.commit()
            
            # Reindex each document
            total_chunks = 0
            for doc in documents:
                if doc.content:
                    chunks = self._create_and_store_chunks(doc.content, doc.id)
                    total_chunks += len(chunks)
            
            return {
                "status": "success",
                "documents_reindexed": len(documents),
                "total_chunks_created": total_chunks
            }
            
        except Exception as e:
            logger.error(f"❌ Error reindexing documents: {e}")
            self.db.rollback()
            return {
                "status": "error",
                "error": str(e)
            }


# Helper function to create RAG service instance
def get_rag_service(db: Session, user_id: int) -> EnhancedRAGService:
    """Factory function to create RAG service instance"""
    return EnhancedRAGService(db=db, user_id=user_id)

# Backward compatibility alias
RAGService = EnhancedRAGService
